# -*- coding: utf-8 -*-
"""
CSP/算法题复盘 Word/PDF 生成器 v12

修复重点：
1. 图块通用化：不再绑定矩阵题，支持 flow / compare / map / table / generic。
2. 保留兼容：reshape、transpose、index 仍可用，但只是通用图块的特例。
3. 代码块恢复标准灰底 + Consolas + 简单 C++ 高亮。
4. 章节编号全中文：一、二、三……十、十一、十二……
5. 自动合并短行，但不破坏代码块、表格、图块、样例。
6. 继续支持 Word + PDF 自动导出。
"""

from __future__ import annotations

import re
import sys
import shutil
import subprocess
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


FONT_CN = "黑体"
FONT_CODE = "Consolas"

COLOR_DARK = "1F2937"
COLOR_BLUE = "1F4E79"
COLOR_BLUE_LIGHT = "D9EAF7"
COLOR_GRAY = "F3F4F6"
COLOR_GREEN_LIGHT = "EAF4EA"
COLOR_YELLOW_LIGHT = "FFF7D6"
COLOR_RED_LIGHT = "FDECEC"


CPP_KEYWORDS = {
    "int", "long", "double", "float", "char", "bool", "void", "return", "for", "while", "if", "else",
    "switch", "case", "break", "continue", "using", "namespace", "std", "vector", "string", "include",
    "cin", "cout", "endl", "const", "auto", "struct", "class", "public", "private", "true", "false",
    "sort", "push_back", "size", "swap"
}


def read_text(path: Path) -> str:
    for enc in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    raise RuntimeError("无法识别 input.txt 编码，请保存为 UTF-8。")


def parse_sections(text: str) -> Dict[str, str]:
    pattern = re.compile(r"^【(.+?)】\s*$", re.M)
    matches = list(pattern.finditer(text))
    sections = {}
    for i, m in enumerate(matches):
        key = m.group(1).strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        sections[key] = text[start:end].strip()
    return sections


def sanitize_filename(s: str) -> str:
    s = re.sub(r"[\\/:*?\"<>|]", "_", s)
    s = s.replace("CSP真题：", "")
    s = s.replace("——完整题解复盘（C++）", "")
    s = s.replace("完整题解复盘（C++）", "")
    s = re.sub(r"\s+", "", s)
    return s.strip("_-— ") or "未命名题目"


def cn_num(n: int) -> str:
    nums = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九"]
    if n <= 10:
        return "十" if n == 10 else nums[n]
    if n < 20:
        return "十" + nums[n - 10]
    tens, ones = divmod(n, 10)
    return nums[tens] + "十" + (nums[ones] if ones else "")


def set_font(run, name=FONT_CN, size: Optional[float] = None, bold: Optional[bool] = None, color: Optional[str] = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(p, before=0, after=4, line=1.15):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def border_cell(cell, color="D1D5DB", sz="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn("w:" + edge))
        if el is None:
            el = OxmlElement("w:" + edge)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)



def paragraph_shading(p, fill: str):
    p_pr = p._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_header_footer(doc, title: str):
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(header_p, after=0)
    hr = header_p.add_run("CSP 真题复盘 · C++")
    set_font(hr, FONT_CN, 9, False, "6B7280")

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(footer_p, after=0)
    r1 = footer_p.add_run("CSP Review Generator · 第 ")
    set_font(r1, FONT_CN, 9, False, "9CA3AF")
    add_page_number(footer_p)
    r2 = footer_p.add_run(" 页")
    set_font(r2, FONT_CN, 9, False, "9CA3AF")


def add_cover_page(doc, title: str, subtitle: str):
    # 轻量封面：标题、关键词标签、生成日期、语言信息
    spacer = doc.add_paragraph()
    set_spacing(spacer, before=30, after=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=8, after=10)
    r = p.add_run(title)
    set_font(r, FONT_CN, 22, True, COLOR_BLUE)

    if subtitle:
        tags = [x.strip() for x in re.split(r"[；;、,\n]+", subtitle) if x.strip()]
        if tags:
            table = doc.add_table(rows=1, cols=len(tags))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, tag in enumerate(tags):
                cell = table.cell(0, i)
                cell.text = tag
                shade_cell(cell, "EAF3FB")
                border_cell(cell, "BBD7EA")
                for pp in cell.paragraphs:
                    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_spacing(pp, after=1)
                    for rr in pp.runs:
                        set_font(rr, FONT_CN, 9.5, False, COLOR_BLUE)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(info, before=18, after=4)
    r2 = info.add_run("C++ · Word / PDF · 自动题解复盘")
    set_font(r2, FONT_CN, 10.5, False, "4B5563")

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(date_p, after=22)
    r3 = date_p.add_run("生成日期：" + datetime.now().strftime("%Y-%m-%d"))
    set_font(r3, FONT_CN, 9.5, False, "9CA3AF")

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(line, before=8, after=0)
    r4 = line.add_run("算法刷题 · 错题整理 · 复盘沉淀")
    set_font(r4, FONT_CN, 11, False, COLOR_BLUE)


def add_text_para(doc, text: str, size=10.5, color=COLOR_DARK, before=0, after=3):
    p = doc.add_paragraph()
    set_spacing(p, before=before, after=after)
    r = p.add_run(text)
    set_font(r, FONT_CN, size, False, color)
    return p


def add_title(doc, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=6, after=6)
    r = p.add_run(title)
    set_font(r, FONT_CN, 20, True, COLOR_BLUE)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p2, after=12)
        r2 = p2.add_run(subtitle)
        set_font(r2, FONT_CN, 10.5, False, "6B7280")


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph()
    set_spacing(p, before=10 if level == 1 else 5, after=6 if level == 1 else 4)
    if level == 1:
        paragraph_shading(p, "EEF6FC")
        p.paragraph_format.left_indent = Cm(0.12)
        r = p.add_run("  " + text)
        set_font(r, FONT_CN, 15, True, COLOR_BLUE)
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "28")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), "1F4E79")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "BBD7EA")
        p_bdr.append(left)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
    else:
        r = p.add_run(text)
        set_font(r, FONT_CN, 12, True, COLOR_BLUE)


def setup_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    doc.styles["Normal"].font.name = FONT_CN
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    doc.styles["Normal"].font.size = Pt(10.5)


def compact_text(text: str) -> str:
    lines = [ln.rstrip() for ln in text.splitlines()]
    out, buf = [], []
    in_code = False
    in_fig = False

    def special(s: str) -> bool:
        st = s.strip()
        if not st:
            return True
        if st.startswith("```") or st.startswith("[[图:") or st == "[[/图]]":
            return True
        if st.startswith("|") and st.endswith("|"):
            return True
        if st.startswith(("•", "-", "*")):
            return True
        if re.match(r"^阶段\s*\d+[：:]", st):
            return True
        if re.match(r"^(版本\s*\d+|样例\s*\d+|输入|输出|解释)[：:]?", st):
            return True
        if re.match(r"^[\d\s\-]+$", st) and any(c.isdigit() for c in st):
            return True
        return False

    def flush():
        nonlocal buf
        if buf:
            s = "".join(buf).strip()
            if s:
                out.append(s)
            buf = []

    for raw in lines:
        st = raw.strip()
        if st.startswith("```"):
            flush()
            out.append(raw)
            in_code = not in_code
            continue
        if in_code:
            out.append(raw)
            continue
        if st.startswith("[[图:"):
            flush()
            in_fig = True
            out.append(raw)
            continue
        if st == "[[/图]]":
            flush()
            in_fig = False
            out.append(raw)
            continue
        if in_fig:
            out.append(raw)
            continue

        if not st:
            flush()
            if out and out[-1] != "":
                out.append("")
            continue

        if special(st):
            flush()
            out.append(raw)
        else:
            buf.append(st)
    flush()

    final = []
    prev_blank = False
    for ln in out:
        blank = not ln.strip()
        if blank and prev_blank:
            continue
        final.append(ln)
        prev_blank = blank
    return "\n".join(final).strip()


def add_code_block(doc, code: str, label: str = "C++"):
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    label_cell = table.cell(0, 0)
    label_cell.text = label
    shade_cell(label_cell, "E5EEF7")
    border_cell(label_cell, "BBD7EA")
    for p in label_cell.paragraphs:
        set_spacing(p, after=0)
        for r in p.runs:
            set_font(r, FONT_CN, 9, False, COLOR_BLUE)

    cell = table.cell(1, 0)
    shade_cell(cell, COLOR_GRAY)
    border_cell(cell, "D1D5DB")
    cell.text = ""

    for line in code.rstrip("\n").splitlines():
        p = cell.add_paragraph()
        set_spacing(p, after=0, line=1.0)
        tokens = re.split(r"(\W+)", line)
        comment_mode = False
        for tok in tokens:
            if tok == "":
                continue
            r = p.add_run(tok)
            if tok.startswith("//") or comment_mode:
                comment_mode = True
                set_font(r, FONT_CODE, 9, False, "6B7280")
            elif tok in CPP_KEYWORDS:
                set_font(r, FONT_CODE, 9, True, "0F4C81")
            elif re.fullmatch(r"\d+", tok):
                set_font(r, FONT_CODE, 9, False, "B45309")
            else:
                set_font(r, FONT_CODE, 9, False, COLOR_DARK)
    try:
        cell._element.remove(cell.paragraphs[0]._element)
    except Exception:
        pass
    doc.add_paragraph()


def add_sample_box(doc, label: str, content: str):
    if not content.strip():
        return
    p = doc.add_paragraph()
    set_spacing(p, before=4, after=2)
    r = p.add_run(label)
    set_font(r, FONT_CN, 10.5, True, COLOR_BLUE)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, COLOR_GRAY)
    border_cell(cell)
    cell.text = ""
    for line in content.strip().splitlines():
        p2 = cell.add_paragraph()
        set_spacing(p2, after=0, line=1.0)
        r2 = p2.add_run(line)
        set_font(r2, FONT_CODE, 9.5, False, COLOR_DARK)
    try:
        cell._element.remove(cell.paragraphs[0]._element)
    except Exception:
        pass


def markdown_table_at(lines, i):
    return i + 1 < len(lines) and lines[i].strip().startswith("|") and re.match(r"^\|\s*[-: ]+\|", lines[i+1].strip())


def parse_markdown_table(lines, i):
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        parts = [p.strip() for p in lines[i].strip().strip("|").split("|")]
        if not all(re.match(r"^[-: ]+$", p) for p in parts):
            rows.append(parts)
        i += 1
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(len(table.columns)):
            cell = table.cell(i, j)
            txt = row[j] if j < len(row) else ""
            cell.text = txt
            fill = COLOR_BLUE_LIGHT if i == 0 else "FFFFFF"
            # 复杂度重点高亮：O(1)、O(N)、O(n×m) 等
            if i > 0 and ("O(" in txt or "O（" in txt):
                fill = COLOR_GREEN_LIGHT if "O(1)" in txt or "O（1）" in txt else COLOR_YELLOW_LIGHT
            shade_cell(cell, fill)
            border_cell(cell)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_spacing(p, after=1)
                for r in p.runs:
                    set_font(r, FONT_CN, 9.5, i == 0, COLOR_DARK)
    doc.add_paragraph()


def parse_figure(lines, i):
    data = {}
    kind = lines[i].strip()[4:-2].strip().lower()
    i += 1
    while i < len(lines):
        st = lines[i].strip()
        if st == "[[/图]]":
            return kind, data, i + 1
        if "：" in st:
            k, v = st.split("：", 1)
            data[k.strip()] = v.strip()
        elif ":" in st:
            k, v = st.split(":", 1)
            data[k.strip()] = v.strip()
        i += 1
    return kind, data, i


def add_generic_figure(doc, kind: str, data: Dict[str, str]):
    title = data.get("标题", "示意图")
    desc = data.get("说明", "")
    p = doc.add_paragraph()
    set_spacing(p, before=6, after=2)
    r = p.add_run("图：" + title)
    set_font(r, FONT_CN, 10.5, True, COLOR_BLUE)
    if desc:
        add_text_para(doc, desc, 9.5, "4B5563", after=3)

    # 通用流程图：步骤：A -> B -> C
    if kind in ("flow", "流程"):
        steps = re.split(r"\s*(?:->|→|，|;|；)\s*", data.get("步骤", "步骤一 -> 步骤二 -> 步骤三"))
        steps = [s for s in steps if s]
        cols = len(steps) * 2 - 1
        table = doc.add_table(rows=1, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, step in enumerate(steps):
            cell = table.cell(0, idx * 2)
            cell.text = step
            if idx * 2 + 1 < cols:
                table.cell(0, idx * 2 + 1).text = "→"

    # 通用对比图
    elif kind in ("compare", "对比", "contrast", "比较"):
        table = doc.add_table(rows=2, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(0, 0).text = data.get("左标题", "方案一")
        table.cell(0, 1).text = ""
        table.cell(0, 2).text = data.get("右标题", "方案二")
        table.cell(1, 0).text = data.get("左内容", data.get("左矩阵", ""))
        table.cell(1, 1).text = "→" if data.get("方向", "→") else ""
        table.cell(1, 2).text = data.get("右内容", data.get("右矩阵", ""))

    # 通用映射图
    elif kind in ("map", "mapping", "映射", "index", "下标"):
        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        pairs = []
        for k in ("左标题", "输入", "矩阵", "右标题", "输出", "下标", "公式", "说明补充"):
            if k in data:
                pairs.append((k, data[k]))
        if not pairs:
            pairs = [("输入", "二维结构"), ("输出", "一维下标"), ("公式", data.get("公式", ""))]
        for k, v in pairs:
            row = table.add_row()
            row.cells[0].text = k
            row.cells[1].text = v.replace(";", "\n")

    # 兼容旧矩阵题图块：reshape / transpose
    elif kind in ("reshape", "transpose"):
        if kind == "reshape":
            table = doc.add_table(rows=2, cols=5)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(0,0).text = data.get("左标题", "原始状态")
            table.cell(0,1).text = ""
            table.cell(0,2).text = data.get("中标题", "中间过程")
            table.cell(0,3).text = ""
            table.cell(0,4).text = data.get("右标题", "目标状态")
            table.cell(1,0).text = data.get("左内容", data.get("左矩阵", "")).replace(";", "\n")
            table.cell(1,1).text = "→"
            table.cell(1,2).text = data.get("中内容", data.get("一维序列", ""))
            table.cell(1,3).text = "→"
            table.cell(1,4).text = data.get("右内容", data.get("右矩阵", "")).replace(";", "\n")
        else:
            table = doc.add_table(rows=2, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(0,0).text = data.get("左标题", "原始状态")
            table.cell(0,1).text = ""
            table.cell(0,2).text = data.get("右标题", "目标状态")
            table.cell(1,0).text = data.get("左内容", data.get("左矩阵", "")).replace(";", "\n")
            table.cell(1,1).text = "→"
            table.cell(1,2).text = data.get("右内容", data.get("右矩阵", "")).replace(";", "\n")

    # 通用表格式图块
    else:
        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for k, v in data.items():
            if k in ("标题", "说明"):
                continue
            row = table.add_row()
            row.cells[0].text = k
            row.cells[1].text = v.replace(";", "\n")

    for i, row in enumerate(table.rows):
        for cell in row.cells:
            shade_cell(cell, COLOR_BLUE_LIGHT if i == 0 else "FFFFFF")
            border_cell(cell, "CBD5E1")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_spacing(p, after=1, line=1.05)
                for r in p.runs:
                    set_font(r, FONT_CN, 9.5, False, COLOR_DARK)
    doc.add_paragraph()


def add_note_box(doc, title, text, fill):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0,0)
    shade_cell(cell, fill)
    border_cell(cell)
    cell.text = ""
    p = cell.paragraphs[0]
    set_spacing(p, before=3, after=3)
    r = p.add_run(title + "：")
    set_font(r, FONT_CN, 10.5, True, COLOR_BLUE)
    r2 = p.add_run(text.replace("\n", " "))
    set_font(r2, FONT_CN, 10.5, False, COLOR_DARK)
    doc.add_paragraph()



def add_dual_note_box(doc, left_title: str, left_text: str, right_title: str, right_text: str):
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = [left_title, right_title]
    contents = [left_text.replace("\n", " "), right_text.replace("\n", " ")]
    fills = [COLOR_RED_LIGHT, COLOR_GREEN_LIGHT]

    for j in range(2):
        hcell = table.cell(0, j)
        hcell.text = headers[j]
        shade_cell(hcell, "E5EEF7")
        border_cell(hcell, "BBD7EA")
        for p in hcell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_spacing(p, after=1)
            for r in p.runs:
                set_font(r, FONT_CN, 10, True, COLOR_BLUE)

        ccell = table.cell(1, j)
        ccell.text = contents[j]
        shade_cell(ccell, fills[j])
        border_cell(ccell)
        for p in ccell.paragraphs:
            set_spacing(p, after=2, line=1.15)
            for r in p.runs:
                set_font(r, FONT_CN, 10, False, COLOR_DARK)
    doc.add_paragraph()


def add_keywords(doc, text):
    kws = [x.strip() for x in re.split(r"[；;、,\n]+", text) if x.strip()]
    if not kws:
        return
    cols = 4
    rows = (len(kws) + cols - 1)//cols
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(rows):
        for j in range(cols):
            idx = i*cols+j
            cell = table.cell(i,j)
            cell.text = kws[idx] if idx < len(kws) else ""
            shade_cell(cell, "F8FAFC" if idx < len(kws) else "FFFFFF")
            border_cell(cell)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_spacing(p, after=1)
                for r in p.runs:
                    set_font(r, FONT_CN, 9.5, False, COLOR_DARK)
    doc.add_paragraph()


def add_rich_text(doc, text, code_label: str = 'C++'):
    text = compact_text(text)
    lines = text.splitlines()
    i = 0
    in_code = False
    code = []
    while i < len(lines):
        line = lines[i]
        st = line.strip()
        if st.startswith("```"):
            if not in_code:
                in_code = True
                code = []
            else:
                in_code = False
                add_code_block(doc, "\n".join(code), code_label)
            i += 1
            continue
        if in_code:
            code.append(line)
            i += 1
            continue
        if st.startswith("[[图:"):
            kind, data, ni = parse_figure(lines, i)
            add_generic_figure(doc, kind, data)
            i = ni
            continue
        if not st:
            i += 1
            continue
        if markdown_table_at(lines, i):
            rows, ni = parse_markdown_table(lines, i)
            add_table(doc, rows)
            i = ni
            continue
        if st.startswith("•"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.45)
            set_spacing(p, after=2)
            r = p.add_run(st)
            set_font(r, FONT_CN, 10.5, False, COLOR_DARK)
            i += 1
            continue
        if re.match(r"^阶段\s*\d+[：:]", st) or re.match(r"^版本\s*\d+[：:]", st):
            p = doc.add_paragraph()
            set_spacing(p, before=5, after=3)
            r = p.add_run(st)
            set_font(r, FONT_CN, 11, True, COLOR_BLUE)
            i += 1
            continue
        add_text_para(doc, st)
        i += 1
    if in_code and code:
        add_code_block(doc, "\n".join(code), code_label)


def add_toc(doc, headings):
    add_heading(doc, "目录")
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, h in enumerate(headings, 1):
        row = table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = h
        for c in row.cells:
            border_cell(c, "FFFFFF", "0")
            for p in c.paragraphs:
                set_spacing(p, after=1)
                for r in p.runs:
                    set_font(r, FONT_CN, 10, False, COLOR_DARK)
    doc.add_paragraph()


def convert_to_pdf(docx_path: Path):
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        print(f"PDF 已生成：{pdf_path}")
        return
    except Exception as e:
        print(f"docx2pdf 转换失败，尝试 LibreOffice：{e}")
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
                           check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            print(f"PDF 已生成：{pdf_path}")
        except Exception as e:
            print(f"LibreOffice PDF 生成失败：{e}")
    else:
        print("未找到 Microsoft Word/docx2pdf 或 LibreOffice，已跳过 PDF 生成。")


def build(input_path: Path):
    raw = read_text(input_path)
    sections = parse_sections(raw)

    # 兼容旧命名
    if "最终AC代码" in sections and "AC代码" not in sections:
        sections["AC代码"] = sections["最终AC代码"]
    for alias in ("代码演进", "历次代码", "中间代码", "过程代码"):
        if alias in sections and "原始代码" not in sections:
            sections["原始代码"] = sections[alias]

    title = sections.get("标题", "CSP真题：未命名——完整题解复盘（C++）")
    subtitle = sections.get("副标题", "")

    output = input_path.parent / f"CSP-{sanitize_filename(title)}-完整题解复盘.docx"

    doc = Document()
    setup_doc(doc)
    add_header_footer(doc, title)
    # add_cover_page(doc, title, subtitle)
    # doc.add_page_break()
    add_title(doc, title, subtitle)

    ordered = [
        "题目背景", "完整题目要求", "输入格式", "输出格式", "输入输出样例",
        "子任务与限制条件", "题目提示与关键区别", "我的完整思考流程复盘",
        "原始代码", "AC代码", "错误分析", "优化过程",
        "可进一步优化的小细节", "复杂度分析", "题解关键词", "总结"
    ]
    headings_for_toc = [x for x in ordered if x == "输入输出样例" or sections.get(x, "").strip() or x in ("题目背景", "完整题目要求")]
    add_toc(doc, headings_for_toc)

    idx = 1
    def h(name):
        nonlocal idx
        add_heading(doc, f"{cn_num(idx)}、{name}")
        idx += 1

    for sec in ("题目背景", "完整题目要求", "输入格式", "输出格式"):
        if sections.get(sec, "").strip():
            h(sec)
            add_rich_text(doc, sections[sec])

    if any(sections.get(f"样例{n}输入", "").strip() for n in ("1", "2")):
        h("输入输出样例")
        for n in ("1", "2"):
            inp = sections.get(f"样例{n}输入", "").strip()
            out = sections.get(f"样例{n}输出", "").strip()
            exp = sections.get(f"样例{n}解释", "").strip()
            if inp or out or exp:
                add_heading(doc, f"样例 {n}", 2)
                add_sample_box(doc, "输入：", inp)
                add_sample_box(doc, "输出：", out)
                if exp:
                    p = doc.add_paragraph()
                    set_spacing(p, before=4, after=1)
                    r = p.add_run("解释：")
                    set_font(r, FONT_CN, 10.5, True, COLOR_BLUE)
                    add_rich_text(doc, exp)

    combined_error_opt = bool(sections.get("错误分析", "").strip() and sections.get("优化过程", "").strip())

    for sec in [
        "子任务与限制条件", "题目提示与关键区别", "我的完整思考流程复盘",
        "原始代码", "AC代码", "错误分析", "优化过程",
        "可进一步优化的小细节", "复杂度分析", "题解关键词", "总结"
    ]:
        text = sections.get(sec, "").strip()
        if not text:
            continue

        if sec == "错误分析" and combined_error_opt:
            h("错误分析与优化过程")
            add_dual_note_box(
                doc,
                "错误分析",
                compact_text(sections.get("错误分析", "")),
                "优化过程",
                compact_text(sections.get("优化过程", ""))
            )
            continue

        if sec == "优化过程" and combined_error_opt:
            continue

        h(sec)
        if sec == "题解关键词":
            add_keywords(doc, text)
        elif sec == "错误分析":
            add_note_box(doc, sec, compact_text(text), COLOR_RED_LIGHT)
        elif sec == "优化过程":
            add_note_box(doc, sec, compact_text(text), COLOR_GREEN_LIGHT)
        elif sec == "原始代码":
            add_rich_text(doc, text, "C++ · 原始代码")
        elif sec == "AC代码":
            add_rich_text(doc, text, "C++ · AC代码")
        else:
            add_rich_text(doc, text)

    doc.save(output)
    print(f"Word 已生成：{output}")
    convert_to_pdf(output)


def main():
    input_path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("input.txt")
    if not input_path.exists():
        print(f"找不到输入文件：{input_path}")
        return
    build(input_path)


if __name__ == "__main__":
    main()
